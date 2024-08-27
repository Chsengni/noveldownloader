import tkinter as tk
from tkinter import messagebox
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import threading
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH  #设置对其方式
import webbrowser

class BookCrawler:
    def __init__(self, root):
        self.root = root
        self.root.title("夜伴书屋小说下载器（www.ybsws.com）")
        # 设置窗口大小
        window_width = 590
        window_height = 600
        self.root.geometry(f"{window_width}x{window_height}")
        
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        
        # 计算窗口的x和y坐标，使其居中
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        
        # 设置窗口位置
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")
        # 创建框架来放置输入框和按钮
        input_frame = tk.Frame(root)
        input_frame.pack(pady=10, padx=10, fill=tk.X)  # 使输入框和按钮适应窗口宽度

        self.book_id_entry = tk.Entry(input_frame, width=60)
        self.book_id_entry.grid(row=0, column=0, padx=5)  # 输入框放在第一列
        self.book_id_entry.insert(0, "18688")

        self.fetch_button = tk.Button(input_frame, text="获取章节", command=self.start_fetch_chapters)
        self.fetch_button.grid(row=0, column=1, padx=5)  # 获取章节按钮放在第二列

        self.download_button = tk.Button(input_frame, text="下载文档", command=self.start_download_chapters)
        self.download_button.grid(row=0, column=2, padx=5)  # 下载章节按钮放在第三列
        self.download_button.config(state=tk.DISABLED)  # 初始时禁用下载按钮

        self.chapter_listbox = tk.Listbox(root, width=80, height=15)
        self.chapter_listbox.pack(pady=10)
        self.chapter_listbox.bind('<Double-1>', self.on_double_click)
        self.status_label = tk.Label(root, text="", fg="red")
        self.status_label.pack(pady=5)

        # 添加文本框用于显示输出
        self.output_text = tk.Text(root, width=80, height=15, wrap=tk.WORD)
        self.output_text.pack(pady=10)
        self.output_text.config(state=tk.DISABLED)  # 初始时禁用文本框编辑

        self.chapters = []  # 存储章节信息
        self.doc = None  # 用于存储文档对象

    def on_double_click(self, event):
        selected_index = self.chapter_listbox.curselection()
        if not selected_index or selected_index[0] == 0:
            return  # 如果是第一项或没有选择项，则不执行任何操作

        # 获取选中的章节
        selected_chapter = self.chapters[selected_index[0] - 1]  # chapters 的索引从 0 开始
        chapter_name, chapter_url = selected_chapter

        # 打开默认浏览器访问该章节的 URL
        webbrowser.open(chapter_url)

    def start_fetch_chapters(self):
        # 启动线程以避免阻塞主线程
        threading.Thread(target=self.fetch_chapters, daemon=True).start()

    def start_download_chapters(self):
        # 启动线程以避免阻塞主线程
        threading.Thread(target=self.download_chapters, daemon=True).start()

    def fetch_chapters(self):
        self.chapter_listbox.delete(0, tk.END)
        self.status_label.config(text="")
        self.chapters = []  # 清空章节列表
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete(1.0, tk.END)
        book_id = self.book_id_entry.get()

        if not book_id:
            self.status_label.config(text="书号不能为空")
            return

        url = f'https://www.ybsws.com/book/{book_id}/'
        headers = {
            'User-Agent': "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36"
        }

        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            response.encoding = response.apparent_encoding
            soup = BeautifulSoup(response.text, 'html.parser')
            book_name_element = soup.find('h1', class_='book-name')

            # 提取书名文本
            if book_name_element:
                book_name = book_name_element.get_text(strip=True)
                #print("书名:", book_name)
                self.chapter_listbox.insert(tk.END, book_name)
            chapter_div = soup.find('div', id='all-chapter')
            if not chapter_div:
                self.status_label.config(text="章节目录未找到")
                return

            links = chapter_div.find_all('a', href=True, title=True)
            for link in links:
                chapter_name = link.get_text()
                chapter_url = link['href']
                self.chapter_listbox.insert(tk.END, chapter_name)
                self.chapters.append((chapter_name, f"https://www.ybsws.com{chapter_url}"))

            self.download_button.config(state=tk.NORMAL)  # 启用下载按钮

        except requests.RequestException as e:
            self.status_label.config(text=f"请求失败: {e}")

    def fetch_pages(self, chapter_url):
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36'}
        pages_text = []

        try:
            response = requests.get(chapter_url, headers=headers)
            response.raise_for_status()
            response.encoding = response.apparent_encoding
            soup = BeautifulSoup(response.text, 'html.parser')
            
            pagination = soup.find('ul', class_='pagination')

            pages = []
            page_links = pagination.find_all('a', href=True)
            pages.append(chapter_url.split(".html")[0].split("/")[-1]+".html")
            for link in page_links:
                href = link['href']
                if href.endswith('.html'):
                    pages.append(href)

            for page in pages:
                full_url = f"https://www.ybsws.com/book/{self.book_id_entry.get()}/{page}"
                self.output_text.insert(tk.END, f"请求页面链接: {full_url}\n")
                self.output_text.see(tk.END)  # 滚动到最新内容
                response = requests.get(full_url, headers=headers)
                response.raise_for_status()
                response.encoding = response.apparent_encoding
                page_soup = BeautifulSoup(response.text, 'html.parser')
                paragraphs = page_soup.find_all('p')
                for para in paragraphs:
                    text = para.get_text()
                    text = text.replace(
                        "如果版权人认为在本站放置您的作品有损您的利益，请发邮件至，本站确认后将会无条件删除。",
                        ""
                    ).replace(
                        "本站所收录作品、社区话题、书库评论均属其个人行为，不代表本站立场。",
                        ""
                    ).replace(
                        "有能力者，请一定订阅和购买正版书籍支持作者，这样作者才能写出更多更好的书！",
                        ""
                    ).replace("\n","").replace(" ","").replace("\r","").replace("\t","")
                    if text:
                        pages_text.append(text)
        except requests.RequestException as e:
            self.status_label.config(text=f"请求失败: {e}")
        
        return pages_text

    def download_chapters(self):
        if not self.chapters:
            self.status_label.config(text="没有章节可下载")
            return
    
        self.doc = Document()  # 初始化文档对象
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete(1.0, tk.END)
        i=0
        for chapter_name, chapter_url in self.chapters:
            self.output_text.insert(tk.END, f"获取章节: {chapter_name}\n")
            self.output_text.see(tk.END)  # 滚动到最新内容 
            if i>0:
                self.doc.add_page_break()
            doc_para = self.doc.add_heading(chapter_name, level=1)  # 添加章节名称作为一级标题
            i=i+1
            try:
                pages_text = self.fetch_pages(chapter_url)
                for text in pages_text:
                    doc_para = self.doc.add_paragraph(text)
                    doc_para.style.font.name = 'Times New Roman'
                    doc_para.style.font.size = Pt(14)  # 四号字体，约等于12pt
                    para_format = doc_para.paragraph_format
                    para_format.first_line_indent =  0
                    para_format.alignment=WD_ALIGN_PARAGRAPH.LEFT#设置为左对
                    para_format.line_spacing = 1.5 # 行距1.5倍，大约18pt
            except Exception as e:
                self.status_label.config(text=f"请求失败: {e}")
                self.save_document()
                return
        
        self.save_document()
        self.status_label.config(text=f"章节内容已保存到 '{self.chapter_listbox.get(0)}.docx'")
        self.output_text.insert(tk.END, "下载完成\n")
        self.output_text.config(state=tk.DISABLED)  # 禁用文本框编辑

    def save_document(self):
        if self.doc:
            try:
                self.doc.save(f'{self.chapter_listbox.get(0)}.docx')
            except Exception as e:
                self.status_label.config(text=f"保存文档失败: {e}")

    def on_closing(self):
        # 确保在关闭窗口时保存文档
        if self.doc:
            self.save_document()
        self.root.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    app = BookCrawler(root)
    root.protocol("WM_DELETE_WINDOW", app.on_closing)  # 处理窗口关闭事件
    root.mainloop()
