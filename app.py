import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import threading
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import webbrowser
from concurrent.futures import ThreadPoolExecutor, as_completed

class BookCrawler:
    def __init__(self, root):
        self.root = root
        self.root.title("夜伴书屋小说下载器（www.ybsws.com）")
        window_width = 590
        window_height = 600
        self.root.geometry(f"{window_width}x{window_height}")
        
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")
        
        input_frame = tk.Frame(root)
        input_frame.pack(pady=10, padx=10, fill=tk.X)

        tk.Label(input_frame, text="书号:").grid(row=0, column=0, padx=5, sticky=tk.W)
        self.book_id_entry = tk.Entry(input_frame, width=30)
        self.book_id_entry.grid(row=0, column=1, padx=5)
        self.book_id_entry.insert(0, "18688")

        tk.Label(input_frame, text="线程池数量:").grid(row=0, column=2, padx=5, sticky=tk.W)
        self.thread_count_combobox = ttk.Combobox(input_frame, values=["1", "2", "4", "8", "16"], width=5)
        self.thread_count_combobox.grid(row=0, column=3, padx=5)
        self.thread_count_combobox.set("4")  # 默认线程池数量为4

        self.fetch_button = tk.Button(input_frame, text="获取章节", command=self.start_fetch_chapters)
        self.fetch_button.grid(row=0, column=4, padx=5)

        self.download_button = tk.Button(input_frame, text="下载文档", command=self.start_download_chapters)
        self.download_button.grid(row=0, column=5, padx=5)
        self.download_button.config(state=tk.DISABLED)

        self.chapter_listbox = tk.Listbox(root, width=80, height=15)
        self.chapter_listbox.pack(pady=10)
        self.chapter_listbox.bind('<Double-1>', self.on_double_click)

        self.status_label = tk.Label(root, text="", fg="red")
        self.status_label.pack(pady=5)

        self.output_text = tk.Text(root, width=80, height=15, wrap=tk.WORD)
        self.output_text.pack(pady=10)
        self.output_text.config(state=tk.DISABLED)

        self.chapters = []
        self.doc = None

        # 绑定窗口关闭事件
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def on_closing(self):
        # 执行保存文档和其他清理操作
        if self.doc:
            self.save_document(self.chapter_listbox.get(0))
        self.root.destroy()

    def on_double_click(self, event):
        selected_index = self.chapter_listbox.curselection()
        if not selected_index or selected_index[0] == 0:
            return
        selected_chapter = self.chapters[selected_index[0] - 1]
        chapter_name, chapter_url = selected_chapter
        webbrowser.open(chapter_url)

    def start_fetch_chapters(self):
        threading.Thread(target=self.fetch_chapters, daemon=True).start()

    def start_download_chapters(self):
        threading.Thread(target=self.download_chapters, daemon=True).start()

    def fetch_chapters(self):
        self.chapter_listbox.delete(0, tk.END)
        self.status_label.config(text="")
        self.chapters = []
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete(1.0, tk.END)
        book_id = self.book_id_entry.get()

        if not book_id:
            self.status_label.config(text="书号不能为空")
            return

        url = f'https://www.ybsws.com/book/{book_id}/'
        headers = {
            'User-Agent': "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36"
        }

        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            response.encoding = response.apparent_encoding
            soup = BeautifulSoup(response.text, 'html.parser')
            book_name_element = soup.find('h1', class_='book-name')

            if book_name_element:
                book_name = book_name_element.get_text(strip=True)
                self.chapter_listbox.insert(tk.END, book_name)
            chapter_div = soup.find('div', id='all-chapter')
            if not chapter_div:
                self.status_label.config(text="章节目录未找到")
                return

            links = chapter_div.find_all('a', href=True, title=True)
            for link in links:
                chapter_name = link.get_text()
                chapter_url = link['href']
                self.chapter_listbox.insert(tk.END, chapter_name)
                self.chapters.append((chapter_name, f"https://www.ybsws.com{chapter_url}"))

            self.download_button.config(state=tk.NORMAL)

        except requests.RequestException as e:
            self.status_label.config(text=f"请求失败: {e}")

    def fetch_pages(self,chapter_name,chapter_url):
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36'}
        pages_text = []
        self.output_text.insert(tk.END, f"获取章节: {chapter_name}\n")
        self.output_text.see(tk.END)
        try:
            response = requests.get(chapter_url, headers=headers)
            response.raise_for_status()
            response.encoding = response.apparent_encoding
            soup = BeautifulSoup(response.text, 'html.parser')
            pagination = soup.find('ul', class_='pagination')
            if pagination:
                pages = [chapter_url.split(".html")[0].split("/")[-1]+".html"]
                page_links = pagination.find_all('a', href=True)
                for link in page_links:
                    href = link['href']
                    if href.endswith('.html'):
                        pages.append(href)
            else:
                pages = [chapter_url.split(".html")[0].split("/")[-1]+".html"]

            for page in pages:
                full_url = f"https://www.ybsws.com/book/{self.book_id_entry.get()}/{page}"
                self.output_text.insert(tk.END, f"请求页面链接: {full_url}\n")
                self.output_text.see(tk.END)
                response = requests.get(full_url, headers=headers)
                response.raise_for_status()
                response.encoding = response.apparent_encoding
                page_soup = BeautifulSoup(response.text, 'html.parser')
                paragraphs = page_soup.find_all('p')
                for para in paragraphs:
                    text = para.get_text()
                    text = text.replace(
                        "如果版权人认为在本站放置您的作品有损您的利益，请发邮件至，本站确认后将会无条件删除。",
                        ""
                    ).replace(
                        "本站所收录作品、社区话题、书库评论均属其个人行为，不代表本站立场。",
                        ""
                    ).replace(
                        "有能力者，请一定订阅和购买正版书籍支持作者，这样作者才能写出更多更好的书！",
                        ""
                    ).replace("\n","").replace(" ","").replace("\r","").replace("\t","")
                    if text:
                        pages_text.append(text)
        except requests.RequestException as e:
            self.status_label.config(text=f"请求失败: {e}")
        
        return pages_text

    def download_chapters(self):
        if not self.chapters:
            self.status_label.config(text="没有章节可下载")
            return

        self.doc = Document()
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete(1.0, tk.END)

        try:
            thread_count = int(self.thread_count_combobox.get())
        except ValueError:
            self.status_label.config(text="线程池数量无效")
            return
        
        if thread_count <= 0:
            self.status_label.config(text="线程池数量必须大于0")
            return
        
        result={}

        with ThreadPoolExecutor(max_workers=thread_count) as executor:
            futures = []
            for chapter_name, chapter_url in self.chapters:
                futures.append(executor.submit(self.fetch_and_write_chapter, chapter_name, chapter_url))
            for future in as_completed(futures):
                chapter_name,pages_text = future.result() 
                result[chapter_name]=pages_text
        i=0
        for chapter in self.chapters:
            if i >1:
                self.doc.add_page_break()
            self.doc.add_heading(chapter[0], level=1)
            i=i+1
            pages_text= result.get(chapter[0], [])
            for page_text in pages_text:
                self.doc.add_paragraph(page_text)

        self.save_document(self.chapter_listbox.get(0))
        self.output_text.config(state=tk.DISABLED)
        self.status_label.config(text="下载完成")

    def fetch_and_write_chapter(self, chapter_name, chapter_url):
        pages_text = self.fetch_pages(chapter_name,chapter_url)
        self.output_text.see(tk.END)
        return chapter_name,pages_text

    def save_document(self,name):
        if self.doc:
            file_path = f"{name}.docx"
            self.doc.save(file_path)

if __name__ == "__main__":
    root = tk.Tk()
    app = BookCrawler(root)
    root.mainloop()
